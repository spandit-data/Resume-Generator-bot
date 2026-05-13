"""DOCX-based resume generator using template and python-docx."""

import logging
import os
import shutil
import subprocess
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

EXPERIENCED_TEMPLATE = Path(__file__).parent / "template" / "resume_template_experience.docx"
FRESHER_TEMPLATE = Path(__file__).parent / "template" / "resume_template_fresher.docx"

FRESHER_KEYWORDS = ["nahi", "nhi", "no", "none", "fresher", "koi nahi", ""]


def is_fresher_data(data: dict) -> bool:
    """Check if user is a fresher based on previous company field."""
    previous_company = data.get("previous_company", "").lower().strip()
    return previous_company in FRESHER_KEYWORDS or not previous_company


def replace_in_doc(doc: Document, old: str, new: str) -> None:
    """Replace placeholder text in document paragraphs."""
    for para in doc.paragraphs:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)


def clear_paragraph_text(para) -> None:
    """Clear all text in a paragraph by clearing runs."""
    for run in para.runs:
        run.text = ""


def generate_docx(user_id: int, data: dict) -> str:
    """Generate resume docx from template using data."""
    is_fresher = is_fresher_data(data)
    template_path = FRESHER_TEMPLATE if is_fresher else EXPERIENCED_TEMPLATE

    docx_path = f"/tmp/resume_{user_id}.docx"
    shutil.copy(template_path, docx_path)

    doc = Document(docx_path)

    # Common replacements for both templates
    replace_in_doc(doc, "{{FULL_NAME}}", data.get("name", "").title())
    replace_in_doc(doc, "{{PHONE}}", data.get("phone", ""))
    replace_in_doc(doc, "{{CITY}}", data.get("city", "").title())
    replace_in_doc(doc, "{{EDUCATION}}", data.get("education", ""))
    replace_in_doc(doc, "{{EDUCATION_YEAR}}", data.get("education_year", ""))

    # Objective line
    job_target = data.get("job_target", "")
    objective = data.get("objective", "")
    if not objective and job_target:
        if is_fresher:
            objective = f"Motivated fresher seeking a {job_target} position to learn and contribute effectively."
        else:
            objective = f"Seeking a {job_target} position where I can contribute my skills and experience effectively."
    replace_in_doc(doc, "{{OBJECTIVE_LINE}}", objective)

    # Skills
    skills = data.get("skills", [])
    for i in range(1, 4):
        skill = skills[i - 1] if i <= len(skills) else ""
        replace_in_doc(doc, f"{{{{SKILL_{i}}}}}", skill)

    # Vehicle - replace placeholder or clear if not provided
    vehicle = data.get("vehicle", "").strip().lower()
    if not vehicle or vehicle in ["", "no", "none", "nahi", "nahi hai", "nhi"]:
        for para in doc.paragraphs:
            if "{{VEHICLE}}" in para.text:
                clear_paragraph_text(para)
    else:
        replace_in_doc(doc, "{{VEHICLE}}", data.get("vehicle", "").title())

    # Only process work experience fields for experienced users
    if not is_fresher:
        replace_in_doc(doc, "{{COMPANY_NAME}}", data.get("previous_company", "").title())
        replace_in_doc(doc, "{{EXPERIENCE_DURATION}}", data.get("experience_duration", ""))
        replace_in_doc(doc, "{{PREVIOUS_ROLE}}", data.get("previous_role", "").title())

        # Work bullets
        work_bullets = data.get("work_bullets", [])
        for i in range(1, 5):
            bullet = work_bullets[i - 1] if i <= len(work_bullets) else ""
            replace_in_doc(doc, f"{{{{WORK_BULLET_{i}}}}}", bullet)

    doc.save(docx_path)
    return docx_path


def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert docx to PDF using LibreOffice."""
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            "/tmp/",
            docx_path,
        ],
        check=True,
    )
    pdf_path = docx_path.replace(".docx", ".pdf")
    return pdf_path


def generate_pdf(user_id: int, data: dict) -> str:
    """Generate a professional PDF resume from docx template."""
    docx_path = generate_docx(user_id, data)
    pdf_path = convert_docx_to_pdf(docx_path)

    logger.info(f"Resume PDF generated: {pdf_path}")
    return pdf_path


def cleanup_files(user_id: int) -> None:
    """Clean up temporary docx and pdf files."""
    for ext in ["docx", "pdf"]:
        path = f"/tmp/resume_{user_id}.{ext}"
        if os.path.exists(path):
            os.remove(path)
            logger.info(f"Cleaned up: {path}")